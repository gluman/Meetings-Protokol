"""Тесты для templates.py — парсер DOCX/txt + генерация промпта."""
import tempfile
from pathlib import Path

from app.templates import (
    CANONICAL_SECTIONS,
    Template,
    TemplateSection,
    generate_prompt,
    parse_docx,
    parse_file,
    parse_text,
)


def test_parse_text_basic():
    text = """Участники
Иванов
Петров

Повестка
Обсуждение квартальных результатов

Вопросы
1. Выручка за Q1
2. Расходы на маркетинг

Решения
1. Увеличить рекламный бюджет на 10%
2. Провести аудит расходов

Открытые вопросы
1. Необходимость нового сотрудника
"""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(text)
        path = Path(f.name)
    sections = parse_text(path)
    names = [s.name for s in sections]
    assert "Участники" in names
    assert "Повестка" in names
    assert "Вопросы" in names
    assert "Решения" in names
    assert "Открытые вопросы" in names
    path.unlink()


def test_parse_text_markdown_headings():
    text = """# Протокол встречи

## Участники
А, Б, В

## Решения
1. Сделать X
"""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(text)
        path = Path(f.name)
    sections = parse_text(path)
    assert any(s.name == "Участники" for s in sections)
    assert any(s.name == "Решения" for s in sections)
    path.unlink()


def test_parse_text_fallback_canonical():
    """Если секций не нашли — должны вернуться канонические (для UI)."""
    text = "Просто какой-то текст без заголовков."
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write(text)
        path = Path(f.name)
    sections = parse_text(path)
    assert sections == []  # парсер ничего не нашёл
    # Но Template ниже добавит канонические
    path.unlink()


def test_parse_file_creates_template():
    text = """Участники
Иванов

Решения
1. X
"""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write(text)
        path = Path(f.name)
    t = parse_file(path, original_filename="example.txt")
    assert isinstance(t, Template)
    assert t.source_format == "txt"
    assert t.source_filename == "example.txt"
    assert len(t.sections) > 0
    assert t.prompt  # автогенерированный промпт не пуст
    assert "Участники" in t.prompt
    assert "Решения" in t.prompt
    path.unlink()


def test_generate_prompt_structure():
    t = Template(
        name="Test",
        source_filename="t.md",
        source_format="md",
        sections=[
            TemplateSection(name="Участники", example_text="Иванов, Петров"),
            TemplateSection(name="Решения", example_text="1. Сделать X"),
        ],
    )
    t.prompt = generate_prompt(t)
    assert "JSON" in t.prompt
    assert "Участники" in t.prompt
    assert "Решения" in t.prompt
    assert "Иванов" in t.prompt  # example text
    assert "Сделать X" in t.prompt


def test_generate_prompt_no_examples():
    t = Template(
        name="Empty",
        source_filename="e.txt",
        source_format="txt",
        sections=[TemplateSection(name=s) for s in CANONICAL_SECTIONS],
    )
    t.prompt = generate_prompt(t)
    for s in CANONICAL_SECTIONS:
        assert s in t.prompt


def test_parse_docx():
    import docx as docxlib
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = Path(f.name)
    doc = docxlib.Document()
    doc.add_heading("Участники", level=1)
    doc.add_paragraph("Иванов")
    doc.add_heading("Решения", level=1)
    doc.add_paragraph("1. Сделать X")
    doc.save(str(path))
    try:
        sections = parse_docx(path)
        names = [s.name for s in sections]
        assert "Участники" in names
        assert "Решения" in names
    finally:
        path.unlink()
